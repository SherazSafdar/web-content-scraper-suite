#import requests
#from bs4 import BeautifulSoup
#from docx import Document

#URL = "https://optimism.io"

#document = Document()
#document.add_heading("Optimism Homepage Content", 0)

#def fetch_page(url):
#    try:
#        response = requests.get(url, timeout=10)
#        if response.status_code == 200 and 'text/html' in response.headers.get('Content-Type', ''):
#            return BeautifulSoup(response.text, 'html.parser')
#    except Exception as e:
#        print(f"Failed to fetch {url} - {e}")
#    return None

#def extract_content(soup, url):
#    document.add_heading(f"Page: {url}", level=1)

#    content_tags = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p'])

#    for tag in content_tags:
#        text = tag.get_text(strip=True)
#        if text and len(text) > 1:
#            if tag.name.startswith('h'):
#                document.add_paragraph(f"{tag.name.upper()}: {text}")
#            elif tag.name == 'p':
#                document.add_paragraph(f"PARA: {text}")

#if __name__ == "__main__":
#    soup = fetch_page(URL)
#    if soup:
#        extract_content(soup, URL)
#        document.save("optimism_homepage.docx")
#        print("✅ Done! Saved to optimism_homepage.docx")



#from selenium import webdriver
#from selenium.webdriver.chrome.options import Options
#from bs4 import BeautifulSoup
#from docx import Document
#from docx.shared import Pt
#import time

# Setup headless Chrome
#options = Options()
#options.add_argument("--headless")
#options.add_argument("--disable-gpu")
#options.add_argument("--window-size=1920x1080")

#driver = webdriver.Chrome(options=options)

# Load Optimism homepage
#url = "https://arbitrum.io/"
#driver.get(url)
#time.sleep(5)  # Wait for JS to fully load

# Parse rendered HTML
#soup = BeautifulSoup(driver.page_source, 'html.parser')
#driver.quit()

# Create Word document
#doc = Document()
#style = doc.styles['Normal']
#font = style.font
#font.name = 'Calibri'
#font.size = Pt(11)

# Add page title
#title = soup.find("title")
#if title:
#    p = doc.add_paragraph()
#    run = p.add_run(f"Title: {title.text.strip()}")
#    run.bold = True
#    run.font.size = Pt(11)

# Extract and add headings and paragraphs
#for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
#    text = tag.get_text(strip=True)
#    if not text:
#        continue

#    para = doc.add_paragraph()
#    run = para.add_run(text)
#    run.font.size = Pt(11)

#    if tag.name.startswith('h'):
#        run.bold = True

# Save the document
#doc.save("Arbitrum_homepage.docx")
#print("✅ Done: 'Arbitrum.docx' saved successfully.")



from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time

# Word document setup
document = Document()
document.add_heading('Arbitrum Homepage Content', 0)

# Configure Chrome to run in headless mode
chrome_options = Options()
chrome_options.add_argument("--headless")
chrome_options.add_argument("--disable-gpu")
chrome_options.add_argument("--window-size=1920,1080")

driver = webdriver.Chrome(options=chrome_options)

def extract_text_with_links(element):
    """Recursively extract text content while preserving hyperlinks"""
    if element.name == 'a':
        href = element.get('href', '')
        text = element.get_text().strip()
        if text and href:
            return f"{text} ({href})"
        return text
    elif element.name in ['br']:
        return '\n'
    else:
        parts = []
        for child in element.children:
            if child.name is None:  # Text node
                parts.append(child.string.strip() if child.string else '')
            else:
                parts.append(extract_text_with_links(child))
        return ''.join(parts)

def scrape_homepage():
    print("🌐 Loading Arbitrum homepage...")
    driver.get("https://arbitrum.io/")
    time.sleep(3)  # Allow page to render
    
    # Handle cookie consent if present
    try:
        cookie_accept = driver.find_element("xpath", "//button[contains(., 'Accept')]")
        cookie_accept.click()
        print("✅ Accepted cookies")
        time.sleep(1)
    except:
        pass
    
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    
    # Find main content container
    main_content = soup.find('main') or soup.body
    
    print("📝 Extracting structured content...")
    
    # Process all content sections
    for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
        try:
            if element.name.startswith('h'):
                # Add heading with appropriate level
                level = int(element.name[1])
                heading = document.add_heading('', level=level)
                run = heading.add_run(element.get_text().strip())
                run.font.size = Pt(24 - (level * 2))
            
            elif element.name == 'p':
                # Process paragraphs with preserved links
                text = extract_text_with_links(element)
                if text.strip():
                    p = document.add_paragraph(text)
            
            elif element.name in ['ul', 'ol']:
                # Process lists and list items
                for item in element.find_all('li', recursive=False):
                    text = extract_text_with_links(item).strip()
                    if text:
                        style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        document.add_paragraph(text, style=style)
        except Exception as e:
            print(f"⚠️ Error processing element: {e}")
            continue

if __name__ == "__main__":
    try:
        scrape_homepage()
        filename = "arbitrum_homepage_content.docx"
        document.save(filename)
        print(f"💾 Saved to {filename}")
        print(f"Total sections extracted: {len(document.paragraphs) + len(document.tables)}")
    finally:
        driver.quit()
        print("🧹 Browser closed")
