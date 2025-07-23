#import requests
#from bs4 import BeautifulSoup
#from docx import Document

#URL = "https://arbitrum.foundation/grants"

#document = Document()
#document.add_heading("Scraped Grant Page Content from arbitrum.foundation/grants", 0)

#def fetch_page(url):
#    try:
#        response = requests.get(url, timeout=10)
#        if response.status_code == 200 and 'text/html' in response.headers.get('Content-Type', ''):
#            return BeautifulSoup(response.text, 'html.parser')
#    except Exception as e:
#        print(f"Failed to fetch {url} - {e}")
#    return None

#def extract_content(soup, url):
#    document.add_heading(f"Page: {url}", level=1)

#    content_tags = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p'])

#    for tag in content_tags:
#        text = tag.get_text(strip=True)
#        if text and len(text) > 1:
#            if tag.name.startswith('h'):
#                document.add_paragraph(f"{tag.name.upper()}: {text}")
#            elif tag.name == 'p':
#                document.add_paragraph(f"PARA: {text}")

#if __name__ == "__main__":
#    soup = fetch_page(URL)
#    if soup:
#        extract_content(soup, URL)
#        document.save("Arbitrum_grants_page.docx")
#        print("✅ Done! Grant content saved to Arbitrum_grants_page.docx")



from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time

# Word document setup
document = Document()
document.add_heading("Scraped Grant Page Content from arbitrum.foundation/grants", 0)

# Configure Chrome to run in headless mode
chrome_options = Options()
chrome_options.add_argument("--headless")
chrome_options.add_argument("--disable-gpu")
chrome_options.add_argument("--window-size=1920,1080")

driver = webdriver.Chrome(options=chrome_options)

def extract_text_with_links(element):
    """Recursively extract text content while preserving hyperlinks"""
    if element.name == 'a':
        href = element.get('href', '')
        text = element.get_text().strip()
        if text and href:
            return f"{text} ({href})"
        return text
    elif element.name in ['br']:
        return '\n'
    else:
        parts = []
        for child in element.children:
            if child.name is None:  # Text node
                parts.append(child.string.strip() if child.string else '')
            else:
                parts.append(extract_text_with_links(child))
        return ''.join(parts)

def scrape_homepage():
    print("🌐 Loading Arbitrum grants...")
    driver.get("https://arbitrum.foundation/grants")
    time.sleep(3)  # Allow page to render
    
    # Handle cookie consent if present
    try:
        cookie_accept = driver.find_element("xpath", "//button[contains(., 'Accept')]")
        cookie_accept.click()
        print("✅ Accepted cookies")
        time.sleep(1)
    except:
        pass
    
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    
    # Find main content container
    main_content = soup.find('main') or soup.body
    
    print("📝 Extracting structured content...")
    
    # Process all content sections
    for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
        try:
            if element.name.startswith('h'):
                # Add heading with appropriate level
                level = int(element.name[1])
                heading = document.add_heading('', level=level)
                run = heading.add_run(element.get_text().strip())
                run.font.size = Pt(24 - (level * 2))
            
            elif element.name == 'p':
                # Process paragraphs with preserved links
                text = extract_text_with_links(element)
                if text.strip():
                    p = document.add_paragraph(text)
            
            elif element.name in ['ul', 'ol']:
                # Process lists and list items
                for item in element.find_all('li', recursive=False):
                    text = extract_text_with_links(item).strip()
                    if text:
                        style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        document.add_paragraph(text, style=style)
        except Exception as e:
            print(f"⚠️ Error processing element: {e}")
            continue

if __name__ == "__main__":
    try:
        scrape_homepage()
        filename = "arbitrum_grantpage_content.docx"
        document.save(filename)
        print(f"💾 Saved to {filename}")
        print(f"Total sections extracted: {len(document.paragraphs) + len(document.tables)}")
    finally:
        driver.quit()
        print("🧹 Browser closed")