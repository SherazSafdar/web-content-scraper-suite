from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from datetime import datetime, timedelta
import time

BASE_URL = "https://optimism.mirror.xyz/"
ONE_YEAR_AGO = datetime.now() - timedelta(days=365)

# Word doc setup
document = Document()
document.add_heading("OPTIMISM Blog Posts (Last 1 Year)", 0)

# Setup Chrome driver (headless)
options = Options()
# options.add_argument("--headless")
driver = webdriver.Chrome(options=options)

def get_scraped_data():
    try:
        global driver, document, BASE_URL
        print("🔍 Loading blog page...")
        driver.get(BASE_URL)
        time.sleep(5)

        # Scroll to load more blogs (if lazy-loaded)
        for _ in range(5):
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)

        soup = BeautifulSoup(driver.page_source, 'html.parser')
        for post in soup.select(".bc5nci45b.bc5nci4u4.bc5nci1aj"):
            try:
                post_url= post.attrs.get("href")
                driver.get(post_url)
                post = BeautifulSoup(driver.page_source, 'html.parser')
                main_heading = post.select_one("._1sjywpl0._1sjywpl1._1q5u8e43._1q5u8e42.bc5nci23u.bc5nci2bm.bc5nci53i.bc5nci4ow.bc5nci316").text
                blog_post = post.select_one(".css-72ne1l > div") # that element can have multiple text or heading or stucture

                """
                docx format:
                url_blog_post post 1
                main heading post 1
                blogpost content post 1

                url_blog_post_2
                main heading post 2
                blogpost content post 2
                
                """

                # Add URL and main heading
                document.add_paragraph(f"URL: {post_url}")
                document.add_heading(main_heading, level=1)

                # Preserve structural elements
                for el in blog_post.find_all(recursive=False):
                    tag = el.name
                    text = el.get_text(strip=True)

                    if not text:
                        continue

                    if tag == "h1":
                        document.add_heading(text, level=2)
                    elif tag == "h2":
                        document.add_heading(text, level=3)
                    elif tag == "h3":
                        document.add_heading(text, level=4)
                    elif tag == "p":
                        document.add_paragraph(text)
                    elif tag == "ul":
                        for li in el.find_all("li"):
                            li_text = li.get_text(strip=True)
                            if li_text:
                                document.add_paragraph(li_text, style="List Bullet")
                    elif tag == "ol":
                        for li in el.find_all("li"):
                            li_text = li.get_text(strip=True)
                            if li_text:
                                document.add_paragraph(li_text, style="List Number")
                    else:
                        document.add_paragraph(text)

                document.add_paragraph("")
                print(f"✅ Saved: {main_heading}")
            except:
                print(f"skipped website url {post_url}")
                driver = webdriver.Chrome(options=options)
    except:
        a = 1

if __name__ == "__main__":
    try:
        get_scraped_data()

        document.save("optimism_1year.docx")
        print("📄 Saved to optimism_blog_1year.docx")

    finally:
        driver.quit()
