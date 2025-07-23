import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from docx import Document

BASE_URL = "https://www.near.org"
visited_urls = set()
document = Document()
document.add_heading("Scraped Content from near.org", 0)

def fetch_page(url):
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200 and 'text/html' in response.headers.get('Content-Type', ''):
            return BeautifulSoup(response.text, 'html.parser')
    except Exception as e:
        print(f"Failed to fetch {url} - {e}")
    return None

def extract_content(soup, url):
    document.add_heading(f"Page: {url}", level=1)
    
    # Get all content in the order it appears: headings + paragraphs
    content_tags = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p'])

    for tag in content_tags:
        text = tag.get_text(strip=True)
        if text and len(text) > 1:
            if tag.name.startswith('h'):
                document.add_paragraph(f"{tag.name.upper()}: {text}")
            elif tag.name == 'p':
                document.add_paragraph(f"PARA: {text}")


def is_valid_internal_link(link):
    if not link or link.startswith("#") or link.startswith("mailto:"):
        return False
    parsed = urlparse(link)
    return parsed.netloc in ("", "www.near.org")

def crawl(url):
    if url in visited_urls:
        return
    visited_urls.add(url)

    soup = fetch_page(url)
    if not soup:
        return

    extract_content(soup, url)

    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"]
        full_url = urljoin(BASE_URL, href)

        if "/blog" in full_url or urlparse(full_url).netloc != "www.near.org":
            continue

        crawl(full_url)

if __name__ == "__main__":
    crawl(BASE_URL)
    document.save("near_org_scraped_data.docx")
    print("✅ Done! Saved to near_org_scraped_data.docx")
