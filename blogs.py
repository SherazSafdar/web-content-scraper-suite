from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
import time

# Set the base blog homepage URL
BASE_URL = "https://blog.arbitrum.io/"

# Prepare Word document
document = Document()
document.add_heading("Arbitrum Blog Posts", 0)

# Setup Chrome driver
options = Options()
# options.add_argument("--headless")  # Uncomment this if you want headless mode
driver = webdriver.Chrome(options=options)

def get_scraped_data():
    global driver, document, BASE_URL
    try:
        print("🔍 Loading Arbitrum blog page...")
        driver.get(BASE_URL)
        time.sleep(5)

        # 🔁 Scroll to load more blogs (lazy-loading)
        for i in range(6):  # Adjust number if not all blogs load
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)

        # Parse the full blog list page after scrolling
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        blog_buttons = soup.select("button.gh-card.post")
        print(f"🧾 Found {len(blog_buttons)} blog posts.")

        for btn in blog_buttons:
            try:
                onclick = btn.get("onclick", "")
                if "window.location=" in onclick:
                    relative_url = onclick.split("'")[1]
                    post_url = urljoin(BASE_URL, relative_url)

                    print(f"🔗 Visiting: {post_url}")
                    driver.get(post_url)
                    time.sleep(2)

                    post_soup = BeautifulSoup(driver.page_source, "html.parser")

                    # Get blog title
                    title_element = post_soup.select_one("h1.gh-article-title")
                    title = title_element.get_text(strip=True) if title_element else "Untitled"

                    # Get blog content
                    content_div = post_soup.select_one(".gh-content")
                    if not content_div:
                        print(f"⚠️ Skipped (no content): {post_url}")
                        continue

                    # Write to DOCX
                    document.add_paragraph(f"URL: {post_url}")
                    document.add_heading(title, level=1)

                    for el in content_div.find_all(recursive=False):
                        tag = el.name
                        text = el.get_text(strip=True)

                        if not text:
                            continue

                        if tag == "h2":
                            document.add_heading(text, level=2)
                        elif tag == "h3":
                            document.add_heading(text, level=3)
                        elif tag == "p":
                            document.add_paragraph(text)
                        elif tag == "ul":
                            for li in el.find_all("li"):
                                li_text = li.get_text(strip=True)
                                if li_text:
                                    document.add_paragraph(li_text, style="List Bullet")
                        elif tag == "ol":
                            for li in el.find_all("li"):
                                li_text = li.get_text(strip=True)
                                if li_text:
                                    document.add_paragraph(li_text, style="List Number")
                        else:
                            document.add_paragraph(text)

                    document.add_paragraph("")  # Blank line between posts
                    print(f"✅ Saved: {title}")
            except Exception as e:
                print(f"⚠️ Error processing post: {e}")
                continue

    except Exception as e:
        print(f"❌ Failed to load or scroll blog list: {e}")

if __name__ == "__main__":
    try:
        get_scraped_data()
        document.save("arbitrum_blog_posts.docx")
        print("📄 All blog posts saved to arbitrum_blog_posts.docx")
    finally:
        driver.quit()
