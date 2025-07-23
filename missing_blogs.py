from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import time

# ✅ Paste ONE blog post URL here
BLOG_URL = "https://www.near.org/blog/regulation-alone-will-not-save-us-from-big-tech"  # Replace this each time

# === Setup Word Document ===
document = Document()
document.add_heading("Single Blog Scraped from NEAR.org", 0)

# === Setup Chrome Headless Driver ===
options = Options()
options.add_argument("--headless")
driver = webdriver.Chrome(options=options)

def scrape_blog(url):
    driver.get(url)
    time.sleep(2)

    soup = BeautifulSoup(driver.page_source, "html.parser")
    print(f"\n🔗 Scraping blog: {url}")

    # === Extract Title ===
    h1 = soup.find("h1")
    title = h1.get_text(strip=True) if h1 else "No Title Found"
    print("📌 Title:", title)

    # === Extract Date ===
    publish_date = None
    time_tag = soup.find("time")
    if time_tag and time_tag.has_attr("datetime"):
        try:
            publish_date = datetime.fromisoformat(time_tag["datetime"].split("T")[0])
        except:
            publish_date = None

    date_str = publish_date.strftime("%B %d, %Y") if publish_date else "Unknown Date"
    print("🗓 Date:", date_str)

    # === Write to DOCX ===
    heading = document.add_paragraph()
    run = heading.add_run(f"Blog URL: {url}\nTITLE: {title}\nPUBLISHED: {date_str}")
    run.bold = True

    document.add_paragraph()  # Blank line

    # === Paragraphs ===
    for p in soup.find_all("p"):
        text = p.get_text(strip=True)
        if text and len(text) > 1:
            document.add_paragraph(text)

    print("✅ Done. Saving file...")

    # === Save with Title or Date in Name ===
    safe_title = title.replace(" ", "_").replace(":", "").replace("/", "")
    filename = f"{safe_title[:40]}_blog.docx"
    document.save(filename)
    print(f"📄 Saved to: {filename}")

if __name__ == "__main__":
    try:
        scrape_blog(BLOG_URL)
    finally:
        driver.quit()
